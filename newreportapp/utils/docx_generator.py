
from docx import Document
from docx.shared import Cm
from docx.shared import RGBColor
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.conf import settings
import os, re


class ReportDocx:
    def __init__(self, filepath, header_report_image):
        """
        Inicializa o documento com cabeçalho configurado.
        :param filepath: Caminho para salvar o arquivo.
        :param header_report_image: Caminho da imagem do cabeçalho.
        """
        try:
            doc_template = os.path.join(settings.BASE_DIR, 'newreportapp' ,'static', 'docx', 'mydoc.docx')
            self.document = Document(doc_template)
        except Exception as e:
            print(f'Erro ao abrir o documento: {e}')
            doc_template = Document()

        self.filepath = filepath

        if self.document.paragraphs and not self.document.paragraphs[0].text.strip():
            p = self.document.paragraphs[0]._element
            p.getparent().remove(p)


        # Configura margens e cabeçalho
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            # section.left_margin = Cm(3)  # Essas duas linhas vão ficar comentadas caso eu queira usar depois.
            # section.right_margin = Cm(2)

            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = header_paragraph.add_run()
            header_run.add_picture(header_report_image, width=Cm(14))
            header_run.space_after = Cm(0.5)

    def clearAllDoc(self):
        """
        Limpa todo o conteúdo do documento, incluindo cabeçalhos e rodapés, deixando-o em branco.
        """
        for paragraph in self.document.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        for table in self.document.tables:
            t_element = table._element
            t_element.getparent().remove(t_element)



    def generateTitle0(self, text):
        """
        Adiciona um título de nível 0 ao documento.
        :param text: Texto do título.
        """
        # Adiciona o parágrafo com estilo Heading 1
        paragraph = self.document.add_paragraph(text)
        
        # Ajusta a formatação do texto
        run = paragraph.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da fonte como preta
        
        # Centraliza o parágrafo
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Define os espaçamentos superior e inferior
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Cm(1)  
        paragraph_format.space_after = Cm(0.5) 



    def generatePreamble(self, text):
        """
        Adiciona um preâmbulo ao documento.
        :param text: Texto do preâmbulo.
        """
        paragraph = self.document.add_paragraph(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = paragraph.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(11)



    def generateTitle1(self, text):
        """
        Adiciona um título de nível 1 ao documento.
        :param text: Texto do título.
        """
        paragraph = self.document.add_paragraph(text, style="Heading 1")
        paragraph.paragraph_format.left_indent = Cm(2)
        run = paragraph.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Cm(1)  
        paragraph_format.space_after = Cm(0.5)



    def generateTitle2(self, text):
        """
        Adiciona um título de nível 2 ao documento.
        :param text: Texto do título.
        """
        paragraph = self.document.add_paragraph(text, style="Heading 2")
        paragraph.paragraph_format.left_indent = Cm(3)
        run = paragraph.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Cm(1)  
        paragraph_format.space_after = Cm(0.5)



    def generateParagraph1(self, text):
        """
        Adiciona um parágrafo justificado sem recuo ao documento.
        :param text: Texto do parágrafo.
        """
        text = text.replace('\r', '').strip()  # fiz isso e resolveu o problema

        paragraphs = text.split('\n')
        for fragments in paragraphs:
            paragraph = self.document.add_paragraph(fragments)
            paragraph.paragraph_format.left_indent = Cm(1)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = paragraph.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(12)



    def generateParagraph2(self, text):
        """
        Adiciona um parágrafo justificado com recuo ao documento.
        :param text: Texto do parágrafo.
        """
        text = text.replace('\r', '').strip()
        paragraphs = text.split('\n')
        for fragments in paragraphs:
            paragraph = self.document.add_paragraph(fragments, style="Descrição")
            paragraph.paragraph_format.left_indent = Cm(2)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = paragraph.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(12)
    
    

    def generateImage(self, url_img, legend_text):
        """
        Adiciona uma imagem ao documento e insere uma legenda numerada.
        :param url_img: Caminho para o arquivo da imagem.
        :param legend_text: Texto da legenda da imagem.
        """
        try:
            # Adiciona a imagem centralizada com largura de 12 cm
            paragraph = self.document.add_paragraph(style="Fotografia")  # Aplicando o estilo "Fotografia"
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_picture(url_img, width=Cm(12))

            # Adiciona a legenda associada à imagem com o estilo "Legenda"
            legend_paragraph = self.document.add_paragraph(legend_text, style='Legenda')  # Passando o texto da legenda
            legend_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centralizando a legenda
        except Exception as e:
            raise IOError(f"Erro ao adicionar a imagem: {e}")


    def generateLegend(self, text):
        """
        Adiciona uma legenda ao documento.
        :param text: Texto da legenda.
        """
        # Adiciona o texto como um parágrafo com o estilo "Caption"
        paragraph = self.document.add_paragraph(text, style="Legenda")
        
        # Configura o alinhamento e ajustes adicionais
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Cm(1)

        # Ajusta propriedades específicas de estilo, caso necessário
        run = paragraph.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.italic = True



    def keyAndValue(self, key, value):
        """
        Adiciona uma linha com chave em negrito e valor na mesma linha ao documento.
        :param key: Texto da chave (em negrito).
        :param value: Texto do valor (sem negrito).
        """
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinhamento à esquerda

        # Configura a chave
        run_key = paragraph.add_run(f"{key}: ")
        # run_key.font.bold = True
        run_key.font.name = "Times New Roman"
        run_key.font.size = Pt(12)

        # Configura o valor
        run_value = paragraph.add_run(value)
        run_value.font.bold = False
        run_value.font.name = "Arial"
        run_value.font.size = Pt(12)



    def generateFooter(self, pretext, text):
        """
        Gera um rodapé exibido em todas as páginas do documento:
        - Primeira linha: texto fornecido.
        - Segunda linha: "Página X de Y".
        Os textos são alinhados à direita e formatados em Times New Roman, tamanho 12.
        """
        # Acessa o rodapé da primeira seção
        section = self.document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Adiciona o texto fixo na primeira linha
        run_text = paragraph.add_run(pretext)
        run_text.font.name = "Times New Roman"
        run_text.font.size = Pt(10)

        run_text = paragraph.add_run('  |  ')
        run_text.font.name = "Times New Roman"
        run_text.font.size = Pt(14)

        run_text = paragraph.add_run(text)  # Quebra de linha após o texto
        run_text.font.name = "Times New Roman"
        run_text.font.size = Pt(10)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Cm(0.5)  
        # paragraph_format.space_after = Cm(0.5)

        run_text = paragraph.add_run('\n')

        # Adiciona a segunda linha com "Página X de Y"
        # "Página "
        run_page_text = paragraph.add_run("Página ")
        run_page_text.font.name = "Times New Roman"
        run_page_text.font.size = Pt(10)

        # Campo dinâmico para a página atual (X)
        this_page = OxmlElement("w:fldSimple")
        this_page.set(qn("w:instr"), "PAGE")  # Campo do Word para número da página atual
        run_this_page = OxmlElement("w:r")
        text_this_page = OxmlElement("w:t")
        text_this_page.text = ""  # Será preenchido pelo Word
        run_this_page.append(text_this_page)
        this_page.append(run_this_page)
        paragraph._p.append(this_page)

        # " de "
        run_of_text = paragraph.add_run(" de ")
        run_of_text.font.name = "Times New Roman"
        run_of_text.font.size = Pt(10)

        # Campo dinâmico para o total de páginas (Y)
        total_pages = OxmlElement("w:fldSimple")
        total_pages.set(qn("w:instr"), "NUMPAGES")  # Campo do Word para total de páginas
        run_total_pages = OxmlElement("w:r")
        text_total_pages = OxmlElement("w:t")
        text_total_pages.text = ""  # Será preenchido pelo Word
        run_total_pages.append(text_total_pages)
        total_pages.append(run_total_pages)
        paragraph._p.append(total_pages)


    def generateSignature(self, expert, text):
        """
        Adiciona uma assinatura ao documento com o formato especificado.
        :param expert: Nome do perito.
        :param text: Texto adicional.
        """
        # Parágrafo para "Assinado digitalmente"
        paragraph = self.document.add_paragraph("Assinado digitalmente")
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = paragraph.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        # Parágrafo para o nome do perito
        paragraph_expert = self.document.add_paragraph(expert)
        paragraph_expert.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run_expert = paragraph_expert.runs[0]
        run_expert.font.name = "Times New Roman"
        run_expert.font.size = Pt(12)
        run_expert.font.italic = True
        paragraph_format = paragraph_expert.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        # Parágrafo para o texto adicional
        paragraph_text = self.document.add_paragraph(text)
        paragraph_text.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run_text = paragraph_text.runs[0]
        run_text.font.name = "Times New Roman"
        run_text.font.size = Pt(12)
        run_text.font.italic = True
        paragraph_format = paragraph_text.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)


    def generateFileName(self, report_number, designation_date):
        sanitized_report_number = re.sub(r'[^\w\s]', '', report_number).replace(' ', '_')
        year = designation_date.year
        return f"{sanitized_report_number}_{year}$.docx"        


    def saveDoc(self):
        """
        Salva o documento no caminho especificado.
        """
        try:
            self.document.save(self.filepath)
        except Exception as e:
            raise IOError(f"Failed to save the document: {e}")

        return self.filepath

